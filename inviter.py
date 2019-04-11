#! python3.6
# inviter.py - create custom invitations from a list of names

import docx
from docx.enum.text import WD_BREAK

# Create a list of the guests from txt file.
guests = open(r'E:\Scripts\Automate\automate_online-materials\guests.txt').readlines()

doc = docx.Document(r'E:\Scripts\Automate\invitations.docx')

# Loop through the guest list and create a custom invite for each.
for guest in guests:
    paraObj1 = doc.add_paragraph('It would be a pleasure to have the company of', 'PyStyle1')
    paraObj2 = doc.add_paragraph(guest, 'PyStyle2')
    paraObj3 = doc.add_paragraph('at 11010 Memory Lane on the Evening of', 'PyStyle1')
    paraObj4 = doc.add_paragraph('April 1st', 'PyStyle3')
    paraObj5 = doc.add_paragraph('at 7 o\' clock', 'PyStyle1')
    # Skip adding a page for the final guest on the list
    if guest != guests[-1]:
        paraObj5.runs[0].add_break(WD_BREAK.PAGE)

# Save document.
doc.save(r'E:\Scripts\Automate\invitations.docx')